"""
===========================================================================
  增量修复脚本 — 不重写正文，仅格式刷表
  1. 读取 main.docx（最完整初稿），保持所有正文文本不变
  2. 全局格式：TNR 10pt，行距最小值 12pt
  3. 替换 Table 1（Scenario）和 Table 2（MSE）内容为已验证数据 + OMML 公式
  4. 所有表格重绘为三线表
===========================================================================
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
import sys, os, re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from latex_to_omml import set_cell_with_math, latex_to_omml, build_cell_paragraph, parse_latex_segments

# ── Constants ──
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

SOURCE = "/Users/wangyaoping/ml-inference-paper/paper/main.docx"
OUTPUT = "/Users/wangyaoping/Desktop/paper_final.docx"

# ── Verified table data ──

TABLE1_SCENARIO = [
    ("Linear",          "$y = X\\beta + \\varepsilon$, $\\beta = (2, -1.5, 0.8, 0, 0)$", "Baseline: OLS optimal"),
    ("Semiparametric",  "$y = X\\beta + 0.3\\sin(X_1) + \\varepsilon$",                   "Mild misspecification"),
    ("Nonlinear",       "$y = \\sin X_1 + \\log(1+|X_2|) + X_3 X_4 + \\varepsilon$",      "Strong nonlinearity"),
    ("High-dim sparse", "$p=100$, 5 non-zero coefficients",                                "Selection + estimation"),
]

TABLE2_MSE = [
    ("highdim", "dnn",      "3.3050", "0.6942"),
    ("highdim", "lasso",    "1.0735", "0.1510"),
    ("highdim", "lightgbm", "3.5685", "0.8672"),
    ("highdim", "mlp",      "7.9618", "1.8603"),
    ("highdim", "ols",      "1.3985", "0.1913"),
    ("highdim", "rf",       "7.1531", "2.0531"),
    ("highdim", "ridge",    "1.3938", "0.1902"),
    ("highdim", "xgboost",  "6.3345", "1.7439"),
    ("linear", "dnn",       "1.2156", "0.1572"),
    ("linear", "lasso",     "1.0516", "0.1234"),
    ("linear", "lightgbm",  "1.6059", "0.2051"),
    ("linear", "mlp",       "1.1082", "0.1286"),
    ("linear", "ols",       "1.0336", "0.1224"),
    ("linear", "rf",        "1.5009", "0.2001"),
    ("linear", "ridge",     "1.0333", "0.1224"),
    ("linear", "xgboost",   "1.6309", "0.2179"),
    ("nonlinear", "dnn",    "1.3440", "0.1551"),
    ("nonlinear", "lasso",  "2.1548", "0.3204"),
    ("nonlinear", "lightgbm","1.5530", "0.2196"),
    ("nonlinear", "mlp",    "1.2196", "0.1528"),
    ("nonlinear", "ols",    "2.1912", "0.3295"),
    ("nonlinear", "rf",     "1.7650", "0.3085"),
    ("nonlinear", "ridge",  "2.1908", "0.3293"),
    ("nonlinear", "xgboost","1.7451", "0.2731"),
    ("semiparametric", "dnn",    "1.2080", "0.1382"),
    ("semiparametric", "lasso",  "1.0555", "0.1270"),
    ("semiparametric", "lightgbm","1.6102", "0.1947"),
    ("semiparametric", "mlp",    "1.0853", "0.1223"),
    ("semiparametric", "ols",    "1.0232", "0.1142"),
    ("semiparametric", "rf",     "1.5374", "0.2077"),
    ("semiparametric", "ridge",  "1.0234", "0.1145"),
    ("semiparametric", "xgboost","1.6723", "0.2090"),
]

TABLE3_HOMECREDIT = [
    ("ols",      "0.0698",  "0.6268"),
    ("ridge",    "0.0698",  "0.1277"),
    ("lasso",    "0.0739",  "0.1835"),
    ("rf",       "0.0691",  "78.1705"),
    ("xgboost",  "0.0683",  "1.7062"),
    ("lightgbm", "0.0681",  "1.8957"),
    ("mlp",      "0.0716",  "31.9594"),
    ("dnn",      "0.0717",  "269.5719"),
]

TABLE4_PIMA = [
    ("ols",      "2821.75"),
    ("ridge",    "2819.98"),
    ("lasso",    "2817.09"),
    ("rf",       "2837.26"),
    ("xgboost",  "3465.29"),
    ("lightgbm", "3322.81"),
    ("mlp",      "6168.69"),
    ("dnn",      "2959.47"),
]

TABLE5_COMPUTE = [
    ("dnn",      "0.45"),
    ("lasso",    "0.00"),
    ("lightgbm", "0.20"),
    ("mlp",      "0.11"),
    ("ols",      "0.00"),
    ("rf",       "0.16"),
    ("ridge",    "0.00"),
    ("xgboost",  "0.18"),
]


# ══════════════════════════════════════════════════════════════
# Formatting helpers
# ══════════════════════════════════════════════════════════════

def set_run_font(run):
    """Force TNR 10pt on a run."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, rPr)
    # Font name
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    # Font size
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '20')  # 10pt
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), '20')


def set_paragraph_spacing(paragraph):
    """Set line spacing to AT_LEAST 12pt, zero space before/after."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        paragraph._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '240')        # 12pt in twips (240 = 12pt)
    spacing.set(qn('w:lineRule'), 'atLeast')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')


def process_paragraph(para):
    """Apply TNR 10pt + line spacing 12pt to a paragraph."""
    # Force font on every run
    for run in para.runs:
        set_run_font(run)
    # Line spacing
    set_paragraph_spacing(para)


def nsdecls(prefix):
    return f'xmlns:{prefix}="{W_NS}"'


# ══════════════════════════════════════════════════════════════
# Table helpers (three-line + content)
# ══════════════════════════════════════════════════════════════

def set_three_line_borders_xml(tbl):
    """Apply three-line table borders at XML level."""
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tbl_pr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def set_cell_text(tc, text, bold=False, align_val="left"):
    """Replace cell content with plain text."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    # Clear all paragraphs
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align_val)
    pPr.append(jc)
    # Line spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '20')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)

    t_elem = OxmlElement('w:t')
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = text
    r.append(t_elem)
    p.append(r)
    tc.append(p)


def set_cell_omml(tc, text, align_val="left"):
    """Set cell text with $...$ converted to OMML math."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    new_p = build_cell_paragraph(text, align_val=align_val)
    # Fix sizing in OMML runs
    for r in new_p.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), '20')

    # Add line spacing to paragraph
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')

    tc.append(new_p)


def set_cell_border_bottom(tc, sz="6"):
    """Add thin bottom border to header cells."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    cell_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(cell_borders)


def set_cell_padding(tc, top=2, bottom=2, left=4, right=4):
    """Set cell padding in twips."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top * 20}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom * 20}" w:type="dxa"/>'
        f'  <w:left w:w="{left * 20}" w:type="dxa"/>'
        f'  <w:right w:w="{right * 20}" w:type="dxa"/>'
        '</w:tcMar>'
    )
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(margins)


def do_vertical_merge(tbl, group_size=8):
    """Merge first column in groups of group_size rows (skip header row 0)."""
    trs = tbl.findall(qn('w:tr'))
    n_groups = (len(trs) - 1) // group_size
    for g in range(n_groups):
        start = 1 + g * group_size
        for r in range(start, start + group_size):
            tc = trs[r].findall(qn('w:tc'))[0]
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            existing_vm = tcPr.find(qn('w:vMerge'))
            if existing_vm is not None:
                tcPr.remove(existing_vm)
            existing_va = tcPr.find(qn('w:vAlign'))
            if existing_va is not None:
                tcPr.remove(existing_va)

            if r == start:
                vm = OxmlElement('w:vMerge')
                vm.set(qn('w:val'), 'restart')
                tcPr.append(vm)
                va = OxmlElement('w:vAlign')
                va.set(qn('w:val'), 'center')
                tcPr.append(va)
            else:
                vm = OxmlElement('w:vMerge')
                vm.set(qn('w:val'), 'continue')
                tcPr.append(vm)
                # Clear text in merged cells
                for p in tc.findall(qn('w:p')):
                    tc.remove(p)
                blank = OxmlElement('w:p')
                blank_r = OxmlElement('w:r')
                blank_t = OxmlElement('w:t')
                blank_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                blank_r.append(blank_t)
                blank.append(blank_r)
                tc.append(blank)


def format_header_row(tr, headers, align="center"):
    """Format first row with bold headers."""
    tcs = tr.findall(qn('w:tc'))
    for ci, text in enumerate(headers):
        if ci < len(tcs):
            set_cell_text(tcs[ci], text, bold=True, align_val=align)
            set_cell_border_bottom(tcs[ci])
            set_cell_padding(tcs[ci])


def format_data_rows(tbl, start_row, data, col_count, aligns=None):
    """Fill data rows with text. aligns: list of alignments per column."""
    trs = tbl.findall(qn('w:tr'))
    if aligns is None:
        aligns = ["left"] + ["center"] * (col_count - 1)
    for ri, row_data in enumerate(data):
        idx = start_row + ri
        if idx >= len(trs):
            break
        tcs = trs[idx].findall(qn('w:tc'))
        for ci in range(min(len(row_data), col_count)):
            if ci < len(tcs):
                set_cell_text(tcs[ci], row_data[ci], align_val=aligns[ci] if ci < len(aligns) else "center")
                set_cell_padding(tcs[ci])


# ══════════════════════════════════════════════════════════════
# Main: incremental fix
# ══════════════════════════════════════════════════════════════

def main():
    print("Opening source: main.docx")
    doc = Document(SOURCE)

    # ── Step 1: Fix all paragraph formatting (TNR 10pt, line spacing) ──
    print(f"\nStep 1: Formatting {len(doc.paragraphs)} paragraphs...")
    for i, para in enumerate(doc.paragraphs):
        process_paragraph(para)
    print("  Done — all paragraphs formatted.")

    # ── Step 2: Fix each table ──
    print(f"\nStep 2: Processing {len(doc.tables)} tables...")

    for ti, table in enumerate(doc.tables):
        tbl = table._tbl
        trs = tbl.findall(qn('w:tr'))
        print(f"  Table {ti}: {len(trs)} rows")

        # Apply three-line border to ALL tables
        set_three_line_borders_xml(tbl)

        if ti == 0:
            # Table 1: Simulation scenarios
            # Header
            tcs = trs[0].findall(qn('w:tc'))
            headers = ["Scenario", "Data generating process", "Purpose"]
            for ci, text in enumerate(headers):
                if ci < len(tcs):
                    set_cell_text(tcs[ci], text, bold=True, align_val="center")
                    set_cell_border_bottom(tcs[ci])
                    set_cell_padding(tcs[ci])

            # Data rows with OMML for column 1
            for ri, (scenario, dgp, purpose) in enumerate(TABLE1_SCENARIO):
                idx = ri + 1
                if idx >= len(trs):
                    break
                tcs = trs[idx].findall(qn('w:tc'))
                if len(tcs) >= 3:
                    set_cell_text(tcs[0], scenario, align_val="left")
                    set_cell_padding(tcs[0])
                    set_cell_omml(tcs[1], dgp, align_val="left")  # $...$ → OMML
                    set_cell_padding(tcs[1])
                    set_cell_text(tcs[2], purpose, align_val="left")
                    set_cell_padding(tcs[2])
            print("    → Three-line + OMML formulas applied.")

        elif ti == 1:
            # Table 2: MSE data
            # Header
            tcs = trs[0].findall(qn('w:tc'))
            headers = ["Scenario", "Method", "MSE", "SD"]
            for ci, text in enumerate(headers):
                if ci < len(tcs):
                    set_cell_text(tcs[ci], text, bold=True, align_val="center")
                    set_cell_border_bottom(tcs[ci])
                    set_cell_padding(tcs[ci])

            # Data rows
            aligns = ["left", "left", "center", "center"]
            for ri, (scenario, method, mse, sd) in enumerate(TABLE2_MSE):
                idx = ri + 1
                if idx >= len(trs):
                    break
                tcs = trs[idx].findall(qn('w:tc'))
                row_data = [scenario, method, mse, sd]
                for ci in range(4):
                    if ci < len(tcs):
                        set_cell_text(tcs[ci], row_data[ci], align_val=aligns[ci])
                        set_cell_padding(tcs[ci])

            # Vertical merge: 4 groups of 8
            do_vertical_merge(tbl, group_size=8)

            # Column widths
            from docx.shared import Inches
            widths = [Inches(1.3), Inches(1.1), Inches(1.0), Inches(1.0)]
            for ri, row in enumerate(table.rows):
                for ci, w in enumerate(widths):
                    if ci < len(row.cells):
                        row.cells[ci].width = w
            print("    → Three-line + vertical merge + data replaced.")

        elif ti == 2:
            # Table 3: Home Credit
            tcs = trs[0].findall(qn('w:tc'))
            headers = ["Method", "Test MSE", "Train Time (s)"]
            for ci, text in enumerate(headers):
                if ci < len(tcs):
                    set_cell_text(tcs[ci], text, bold=True, align_val="center")
                    set_cell_border_bottom(tcs[ci])
                    set_cell_padding(tcs[ci])
            aligns = ["left", "center", "center"]
            for ri, (method, mse, ttime) in enumerate(TABLE3_HOMECREDIT):
                idx = ri + 1
                if idx >= len(trs):
                    break
                tcs = trs[idx].findall(qn('w:tc'))
                row_data = [method, mse, ttime]
                for ci in range(3):
                    if ci < len(tcs):
                        set_cell_text(tcs[ci], row_data[ci], align_val=aligns[ci])
                        set_cell_padding(tcs[ci])
            print("    → Three-line + data replaced.")

        elif ti == 3:
            # Table 4: PIMA Diabetes
            tcs = trs[0].findall(qn('w:tc'))
            headers = ["Method", "Test MSE"]
            for ci, text in enumerate(headers):
                if ci < len(tcs):
                    set_cell_text(tcs[ci], text, bold=True, align_val="center")
                    set_cell_border_bottom(tcs[ci])
                    set_cell_padding(tcs[ci])
            for ri, (method, mse) in enumerate(TABLE4_PIMA):
                idx = ri + 1
                if idx >= len(trs):
                    break
                tcs = trs[idx].findall(qn('w:tc'))
                row_data = [method, mse]
                for ci in range(2):
                    if ci < len(tcs):
                        set_cell_text(tcs[ci], row_data[ci], align_val=["left", "center"][ci])
                        set_cell_padding(tcs[ci])
            print("    → Three-line applied.")

        elif ti == 4:
            # Table 5: Computation cost
            tcs = trs[0].findall(qn('w:tc'))
            headers = ["Method", "Train Time (s)"]
            for ci, text in enumerate(headers):
                if ci < len(tcs):
                    set_cell_text(tcs[ci], text, bold=True, align_val="center")
                    set_cell_border_bottom(tcs[ci])
                    set_cell_padding(tcs[ci])
            for ri, (method, ttime) in enumerate(TABLE5_COMPUTE):
                idx = ri + 1
                if idx >= len(trs):
                    break
                tcs = trs[idx].findall(qn('w:tc'))
                row_data = [method, ttime]
                for ci in range(2):
                    if ci < len(tcs):
                        set_cell_text(tcs[ci], row_data[ci], align_val=["left", "center"][ci])
                        set_cell_padding(tcs[ci])
            print("    → Three-line applied.")

    # ── Step 3: Set margins ──
    sect_pr = doc.sections[0]
    sect_pr.top_margin = Inches(1)
    sect_pr.bottom_margin = Inches(1)
    sect_pr.left_margin = Inches(1)
    sect_pr.right_margin = Inches(1)

    # ── Step 4: Verify ──
    print("\nStep 3: Verification...")
    total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    print(f"  Word count: {total_words}")
    print(f"  Tables: {len(doc.tables)}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")

    # Check font compliance
    non_tnr = 0
    for p in doc.paragraphs:
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ascii_font = rFonts.get(qn('w:ascii'))
                    if ascii_font and ascii_font != FONT_NAME:
                        non_tnr += 1
    print(f"  Non-TNR runs: {non_tnr}")

    # Save
    print(f"\nSaving to: {OUTPUT}")
    doc.save(OUTPUT)
    print("Done!")


if __name__ == "__main__":
    main()
