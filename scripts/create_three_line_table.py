"""
Create an academic three-line table (三线表) in a Word document
using python-docx with direct XML manipulation for borders.

Three-line table rules:
  - Only 3 horizontal lines: top (thick), header-bottom (thin), bottom (thick)
  - NO vertical lines
  - NO internal horizontal lines (between data rows)
  - First column identical values are merged vertically
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── 1. DATA ────────────────────────────────────────────────────────────
# Each tuple: (scenario, method, mse, sd)
# Scenarios repeat for each method (8 methods per scenario)
DATA = [
    # highdim scenario
    ("highdim", "dnn",      "3.3050", "0.6942"),
    ("highdim", "lasso",    "1.0735", "0.1510"),
    ("highdim", "lightgbm", "3.5685", "0.8672"),
    ("highdim", "mlp",      "7.9618", "1.8603"),
    ("highdim", "ols",      "1.3985", "0.1913"),
    ("highdim", "rf",       "7.1531", "2.0531"),
    ("highdim", "ridge",    "1.3938", "0.1902"),
    ("highdim", "xgboost",  "6.3345", "1.7439"),
    # linear scenario
    ("linear", "dnn",       "1.2156", "0.1572"),
    ("linear", "lasso",     "1.0516", "0.1234"),
    ("linear", "lightgbm",  "1.6059", "0.2051"),
    ("linear", "mlp",       "1.1082", "0.1286"),
    ("linear", "ols",       "1.0336", "0.1224"),
    ("linear", "rf",        "1.5009", "0.2001"),
    ("linear", "ridge",     "1.0333", "0.1224"),
    ("linear", "xgboost",   "1.6309", "0.2179"),
    # nonlinear scenario
    ("nonlinear", "dnn",    "1.3440", "0.1551"),
    ("nonlinear", "lasso",  "2.1548", "0.3204"),
    ("nonlinear", "lightgbm","1.5530", "0.2196"),
    ("nonlinear", "mlp",    "1.2196", "0.1528"),
    ("nonlinear", "ols",    "2.1912", "0.3295"),
    ("nonlinear", "rf",     "1.7650", "0.3085"),
    ("nonlinear", "ridge",  "2.1908", "0.3293"),
    ("nonlinear", "xgboost","1.7451", "0.2731"),
    # semiparametric scenario
    ("semiparametric", "dnn",    "1.2080", "0.1382"),
    ("semiparametric", "lasso",  "1.0555", "0.1270"),
    ("semiparametric", "lightgbm","1.6102", "0.1947"),
    ("semiparametric", "mlp",    "1.0853", "0.1223"),
    ("semiparametric", "ols",    "1.0232", "0.1142"),
    ("semiparametric", "rf",     "1.5374", "0.2077"),
    ("semiparametric", "ridge",  "1.0234", "0.1145"),
    ("semiparametric", "xgboost","1.6723", "0.2090"),
]

HEADERS = ["scenario", "method", "MSE", "SD"]
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10.5)  # 五号

# ── 2. CREATE DOCUMENT ──────────────────────────────────────────────────
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = FONT_SIZE

# ── 3. BUILD TABLE ──────────────────────────────────────────────────────
num_rows = len(DATA) + 1  # +1 for header
num_cols = len(HEADERS)
table = doc.add_table(rows=num_rows, cols=num_cols)

# ── 4. THREE-LINE TABLE BORDERS (XML manipulation) ──────────────────────
"""
Word table borders are controlled via w:tblBorders element inside w:tblPr.
We need:
  - w:top:     single line, 1.5pt (thick)   → table top border
  - w:bottom:  single line, 1.5pt (thick)   → table bottom border
  - All others (left, right, insideH, insideV): set to "nil" (no border)

The XML structure we build:
  <w:tblBorders>
    <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
  </w:tblBorders>

Note: w:sz is measured in eighths of a point.
  1.5 pt = 12 eighths
  0.75 pt = 6 eighths
"""
tbl = table._tbl  # Access the underlying XML element
tbl_pr = tbl.find(qn('w:tblPr'))
if tbl_pr is None:
    tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    tbl.insert(0, tbl_pr)

# Create tblBorders element
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    # ── Top border: thick (1.5pt = 12 eighths) ────────────────────
    '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    # ── Left / Right / InsideV: no border ─────────────────────────
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    # ── InsideH: no border between DATA rows ──────────────────────
    #   (The header-bottom line is handled per-cell in step 6)
    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    # ── Bottom border: thick (1.5pt = 12 eighths) ─────────────────
    '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    '</w:tblBorders>'
)
# Remove any existing borders first
existing = tbl_pr.find(qn('w:tblBorders'))
if existing is not None:
    tbl_pr.remove(existing)
tbl_pr.append(borders)

# ── 5. SET CELL PADDING (margins) ──────────────────────────────────────
"""
Cell padding prevents text from touching borders.
We set top=bottom=2pt, left=right=4pt on every cell.
The XML is: <w:tcMar> inside w:tcPr
"""
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is None:
            tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tc_pr)
        # Set cell margins (in twentieths of a point: 2pt = 40, 4pt = 80)
        margins = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            '  <w:top w:w="40" w:type="dxa"/>'
            '  <w:bottom w:w="40" w:type="dxa"/>'
            '  <w:left w:w="80" w:type="dxa"/>'
            '  <w:right w:w="80" w:type="dxa"/>'
            '</w:tcMar>'
        )
        existing_mar = tc_pr.find(qn('w:tcMar'))
        if existing_mar is not None:
            tc_pr.remove(existing_mar)
        tc_pr.append(margins)

# ── 6. HEADER ROW: bottom border (thin, 0.75pt) ────────────────────────
"""
The header-bottom line is NOT a table-level border (which would apply to
every row). Instead, we add a bottom border to EACH CELL in the header row.
This way only the row below the header gets a line.

XML: <w:tcBorders> inside w:tcPr with <w:bottom w:val="single" w:sz="6" ...>
0.75 pt = 6 eighths
"""
for cell in table.rows[0].cells:
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tc_pr)
    cell_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    existing_borders = tc_pr.find(qn('w:tcBorders'))
    if existing_borders is not None:
        tc_pr.remove(existing_borders)
    tc_pr.append(cell_borders)

# ── 7. FILL HEADER ROW ─────────────────────────────────────────────────
for col_idx, header_text in enumerate(HEADERS):
    cell = table.rows[0].cells[col_idx]
    # Clear default paragraph
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(header_text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 8. FILL DATA ROWS ──────────────────────────────────────────────────
for row_idx, (scenario, method, mse, sd) in enumerate(DATA, start=1):
    # Column 0: scenario name → left-aligned (horizontal merge handled later)
    cell0 = table.rows[row_idx].cells[0]
    p0 = cell0.paragraphs[0]
    p0.text = ""
    run0 = p0.add_run(scenario)
    run0.font.name = FONT_NAME
    run0.font.size = FONT_SIZE
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Column 1: method name → left-aligned
    cell1 = table.rows[row_idx].cells[1]
    p1 = cell1.paragraphs[0]
    p1.text = ""
    run1 = p1.add_run(method)
    run1.font.name = FONT_NAME
    run1.font.size = FONT_SIZE
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Column 2: MSE → center-aligned
    cell2 = table.rows[row_idx].cells[2]
    p2 = cell2.paragraphs[0]
    p2.text = ""
    run2 = p2.add_run(mse)
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column 3: SD → center-aligned
    cell3 = table.rows[row_idx].cells[3]
    p3 = cell3.paragraphs[0]
    p3.text = ""
    run3 = p3.add_run(sd)
    run3.font.name = FONT_NAME
    run3.font.size = FONT_SIZE
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 9. MERGE IDENTICAL SCENARIO CELLS IN COLUMN 0 ─────────────────────
"""
Scan column 0 top-to-bottom. For each group of consecutive identical
scenario names, merge the cells vertically using direct XML manipulation.

Word vertical merge XML:
  - First cell:  <w:vMerge w:val="restart"/>  (start of merge)
  - Continuation: <w:vMerge/>                  (continue merge)
  Word hides the content of continuation cells and only displays
  the first cell's content. We also clear continuation cells' text
  to avoid concatenated text in python-docx's .text property.

Column 0 data was filled in step 8 — each cell contains the scenario
name. After merge, only the first cell's text is visible in Word.
"""
row_idx = 1  # Start AFTER header row
while row_idx <= len(DATA):
    current_scenario = DATA[row_idx - 1][0]  # scenario name at this row
    start_row = row_idx
    # Find all consecutive rows with the same scenario
    while row_idx <= len(DATA) and DATA[row_idx - 1][0] == current_scenario:
        row_idx += 1
    end_row = row_idx - 1

    # Only merge if there are 2+ rows in this group
    if end_row > start_row:
        # Table rows: row 0 = header, row 1 = DATA[0], etc.
        for r in range(start_row, end_row + 1):
            tc = table.rows[r].cells[0]._tc
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)

            if r == start_row:
                # First cell: vMerge restart + vAlign center
                vm = OxmlElement('w:vMerge')
                vm.set(qn('w:val'), 'restart')
                va = OxmlElement('w:vAlign')
                va.set(qn('w:val'), 'center')
                existing_vm = tc_pr.find(qn('w:vMerge'))
                if existing_vm is not None:
                    tc_pr.remove(existing_vm)
                existing_va = tc_pr.find(qn('w:vAlign'))
                if existing_va is not None:
                    tc_pr.remove(existing_va)
                tc_pr.append(vm)
                tc_pr.append(va)
            else:
                # Continuation cells: vMerge continue, clear text
                # Note: python-docx save/load normalizes val-less vMerge to
                # val="restart", so we explicitly set val="continue"
                vm = OxmlElement('w:vMerge')
                vm.set(qn('w:val'), 'continue')
                existing_vm = tc_pr.find(qn('w:vMerge'))
                if existing_vm is not None:
                    tc_pr.remove(existing_vm)
                tc_pr.append(vm)
                # Clear paragraph content from continuation cells
                for p in tc.findall(qn('w:p')):
                    tc.remove(p)
                # Add a single empty paragraph
                new_p = OxmlElement('w:p')
                new_r = OxmlElement('w:r')
                new_t = OxmlElement('w:t')
                new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_r.append(new_t)
                new_p.append(new_r)
                tc.append(new_p)

# ── 10. SET COLUMN WIDTHS (OPTIONAL, FOR BETTER APPEARANCE) ──────────
"""
Set approximate widths: scenario column ~1.2in, method ~1.0in,
MSE ~1.0in, SD ~1.0in
"""
widths = [Inches(1.3), Inches(1.1), Inches(1.0), Inches(1.0)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# ── 11. SAVE ──────────────────────────────────────────────────────────
output_path = "/Users/wangyaoping/ml-inference-paper/tables/three_line_table.docx"
doc.save(output_path)
print(f"Three-line table saved to: {output_path}")
print(f"Data rows: {len(DATA)}, Total rows (incl. header): {num_rows}")
