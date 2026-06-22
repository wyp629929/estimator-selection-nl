"""
===========================================================================
  克隆 + 原地手术 — 确保正文零丢失
  1. shutil.copyfile 完整克隆 main.docx → paper_final_perfect.docx
  2. 打开克隆文件，原地修改（不新建文档，不复制文本）
  3. 替换 Table 1 / Table 2 为三线表 + OMML 公式
  4. 全局 TNR 10pt + 固定行距 12pt
  5. 验证文件完整性
===========================================================================
"""

import shutil, os, sys, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from latex_to_omml import build_cell_paragraph

# ── Paths ──
SOURCE = "/Users/wangyaoping/ml-inference-paper/paper/main.docx"
OUTPUT = "/Users/wangyaoping/Desktop/paper_final_perfect.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FONT_NAME = "Times New Roman"

# ── Verified table data ──
TABLE1_DATA = [
    ("Linear",          "$y = X\\beta + \\varepsilon$, $\\beta = (2, -1.5, 0.8, 0, 0)$", "Baseline: OLS optimal"),
    ("Semiparametric",  "$y = X\\beta + 0.3\\sin(X_1) + \\varepsilon$",                   "Mild misspecification"),
    ("Nonlinear",       "$y = \\sin X_1 + \\log(1+|X_2|) + X_3 X_4 + \\varepsilon$",      "Strong nonlinearity"),
    ("High-dim sparse", "$p=100$, 5 non-zero coefficients",                                "Selection + estimation"),
]

TABLE2_MSE = [
    ("highdim", "dnn",      "3.3050", "0.6942"),
    ("highdim", "lasso",    "1.0735", "0.1510"),
    ("highdim", "lightgbm", "3.5685", "0.8672"),
    ("highdim", "mlp",      "7.9618", "1.8603"),
    ("highdim", "ols",      "1.3985", "0.1913"),
    ("highdim", "rf",       "7.1531", "2.0531"),
    ("highdim", "ridge",    "1.3938", "0.1902"),
    ("highdim", "xgboost",  "6.3345", "1.7439"),
    ("linear", "dnn",       "1.2156", "0.1572"),
    ("linear", "lasso",     "1.0516", "0.1234"),
    ("linear", "lightgbm",  "1.6059", "0.2051"),
    ("linear", "mlp",       "1.1082", "0.1286"),
    ("linear", "ols",       "1.0336", "0.1224"),
    ("linear", "rf",        "1.5009", "0.2001"),
    ("linear", "ridge",     "1.0333", "0.1224"),
    ("linear", "xgboost",   "1.6309", "0.2179"),
    ("nonlinear", "dnn",    "1.3440", "0.1551"),
    ("nonlinear", "lasso",  "2.1548", "0.3204"),
    ("nonlinear", "lightgbm","1.5530", "0.2196"),
    ("nonlinear", "mlp",    "1.2196", "0.1528"),
    ("nonlinear", "ols",    "2.1912", "0.3295"),
    ("nonlinear", "rf",     "1.7650", "0.3085"),
    ("nonlinear", "ridge",  "2.1908", "0.3293"),
    ("nonlinear", "xgboost","1.7451", "0.2731"),
    ("semiparametric", "dnn",    "1.2080", "0.1382"),
    ("semiparametric", "lasso",  "1.0555", "0.1270"),
    ("semiparametric", "lightgbm","1.6102", "0.1947"),
    ("semiparametric", "mlp",    "1.0853", "0.1223"),
    ("semiparametric", "ols",    "1.0232", "0.1142"),
    ("semiparametric", "rf",     "1.5374", "0.2077"),
    ("semiparametric", "ridge",  "1.0234", "0.1145"),
    ("semiparametric", "xgboost","1.6723", "0.2090"),
]

TABLE3_HOMECREDIT = [
    ("ols",      "0.0698",  "0.6268"),
    ("ridge",    "0.0698",  "0.1277"),
    ("lasso",    "0.0739",  "0.1835"),
    ("rf",       "0.0691",  "78.1705"),
    ("xgboost",  "0.0683",  "1.7062"),
    ("lightgbm", "0.0681",  "1.8957"),
    ("mlp",      "0.0716",  "31.9594"),
    ("dnn",      "0.0717",  "269.5719"),
]

TABLE4_PIMA = [
    ("ols",      "2821.75"),
    ("ridge",    "2819.98"),
    ("lasso",    "2817.09"),
    ("rf",       "2837.26"),
    ("xgboost",  "3465.29"),
    ("lightgbm", "3322.81"),
    ("mlp",      "6168.69"),
    ("dnn",      "2959.47"),
]

TABLE5_COMPUTE = [
    ("dnn",      "0.45"),
    ("lasso",    "0.00"),
    ("lightgbm", "0.20"),
    ("mlp",      "0.11"),
    ("ols",      "0.00"),
    ("rf",       "0.16"),
    ("ridge",    "0.00"),
    ("xgboost",  "0.18"),
]


def nsdecls(p):
    return f'xmlns:{p}="{W_NS}"'


# ══════════════════════════════════════════════════════════════
# Formatting helpers — operate on XML directly for reliability
# ══════════════════════════════════════════════════════════════

def set_spacing_exactly12(para):
    """Set line spacing: Exactly 12pt, zero before/after."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '240')       # 12pt in twips * 20
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')


def set_run_font_tnr10(run):
    """Force TNR 10pt on a single run at XML level."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    # Size
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '20')
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), '20')


def format_paragraph(para):
    """Apply TNR 10pt + exactly 12pt spacing to a paragraph."""
    set_spacing_exactly12(para)
    for run in para.runs:
        set_run_font_tnr10(run)


# ══════════════════════════════════════════════════════════════
# Three-line table helpers
# ══════════════════════════════════════════════════════════════

def apply_three_line_borders(table):
    """Set three-line table borders (top/bottom 1.5pt, header-bottom 0.75pt)."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def clear_table(table):
    """Remove all rows from a table except keep structure."""
    tbl = table._tbl
    for tr in list(tbl.findall(qn('w:tr'))):
        tbl.remove(tr)


def make_cell_paragraph(text, bold=False, align="center"):
    """Build a w:p element with formatted text."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    # Spacing exactly 12pt
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '20')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)

    t_el = OxmlElement('w:t')
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = text
    r.append(t_el)
    p.append(r)
    return p


def make_cell_paragraph_omml(text, align="left"):
    """Build w:p with $...$ converted to OMML."""
    if '$' not in text:
        return make_cell_paragraph(text, align=align)

    new_p = build_cell_paragraph(text, align_val=align)
    # Fix font sizing in all runs
    for r in new_p.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), '20')
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), '20')

    # Ensure spacing on this paragraph
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')

    return new_p


def set_cell_content(tc, paragraph_element):
    """Clear a tc and set its content to one paragraph."""
    # Remove all existing paragraphs
    for p in tc.findall(qn('w:p')):
        tc._element.remove(p)
    # Also clear any tc content via lxml on _tc
    for p in tc._tc.findall(qn('w:p')):
        tc._tc.remove(p)
    tc._tc.append(paragraph_element)


def _tc_elem(tc):
    """Get the raw oxml tc element from either a python-docx Cell or raw element."""
    if hasattr(tc, '_tc'):
        return tc._tc
    return tc


def add_tc_bottom_border(tc):
    """Add 0.75pt bottom border to header cell."""
    tc_e = _tc_elem(tc)
    tcPr = tc_e.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc_e.insert(0, tcPr)
    # Remove existing tcBorders
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_padding(tc, top=40, bottom=40, left=80, right=80):
    """Set cell padding in twips."""
    tc_e = _tc_elem(tc)
    tcPr = tc_e.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc_e.insert(0, tcPr)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_col_width(tc_elem, width_twips):
    """Set column width on a tc oxml element."""
    tcPr = tc_elem.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc_elem.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def build_table1_xml():
    """Build complete Table 1 (Simulation scenarios) as XML."""
    tbl = OxmlElement('w:tbl')
    # Table properties with three-line borders
    tbl_pr = OxmlElement('w:tblPr')
    borders_elem = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders_elem)
    tbl.append(tbl_pr)

    col_widths = [1440, 5040, 2880]  # twips

    # Header row
    h_tr = OxmlElement('w:tr')
    headers = ["Scenario", "Data generating process", "Purpose"]
    for ci, h in enumerate(headers):
        tc = OxmlElement('w:tc')
        set_col_width(tc, col_widths[ci])
        tc.append(make_cell_paragraph(h, bold=True, align="center"))
        add_tc_bottom_border(tc)
        set_cell_padding(tc)
        h_tr.append(tc)
    tbl.append(h_tr)

    # Data rows
    for scenario, dgp, purpose in TABLE1_DATA:
        tr = OxmlElement('w:tr')
        # Scenario
        tc0 = OxmlElement('w:tc')
        set_col_width(tc0, col_widths[0])
        tc0.append(make_cell_paragraph(scenario, align="left"))
        set_cell_padding(tc0)
        tr.append(tc0)
        # DGP (with OMML)
        tc1 = OxmlElement('w:tc')
        set_col_width(tc1, col_widths[1])
        tc1.append(make_cell_paragraph_omml(dgp, align="left"))
        set_cell_padding(tc1)
        tr.append(tc1)
        # Purpose
        tc2 = OxmlElement('w:tc')
        set_col_width(tc2, col_widths[2])
        tc2.append(make_cell_paragraph(purpose, align="left"))
        set_cell_padding(tc2)
        tr.append(tc2)
        tbl.append(tr)

    return tbl


def build_table2_xml():
    """Build Table 2 (Test MSE) with vertical merge."""
    tbl = OxmlElement('w:tbl')
    # Three-line borders
    tbl_pr = OxmlElement('w:tblPr')
    borders_elem = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders_elem)

    # Table width
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'dxa')
    tbl_pr.append(tblW)
    tbl.append(tbl_pr)

    col_widths = [1872, 1584, 1440, 1440]  # twips: 1.3, 1.1, 1.0, 1.0

    # Header row
    h_tr = OxmlElement('w:tr')
    headers = ["Scenario", "Method", "MSE", "SD"]
    for ci, h in enumerate(headers):
        tc = OxmlElement('w:tc')
        set_col_width(tc, col_widths[ci])
        tc.append(make_cell_paragraph(h, bold=True, align="center"))
        add_tc_bottom_border(tc)
        set_cell_padding(tc)
        h_tr.append(tc)
    tbl.append(h_tr)

    # Data rows with vertical merge on column 0
    group_size = 8
    n_groups = len(TABLE2_MSE) // group_size

    for g in range(n_groups):
        start = g * group_size
        for ri in range(start, start + group_size):
            scenario, method, mse, sd = TABLE2_MSE[ri]
            tr = OxmlElement('w:tr')

            # Column 0 - vertically merged
            tc0 = OxmlElement('w:tc')
            tcPr0 = OxmlElement('w:tcPr')
            tc0.insert(0, tcPr0)
            vm = OxmlElement('w:vMerge')
            if ri == start:
                vm.set(qn('w:val'), 'restart')
                tcPr0.append(vm)
                va = OxmlElement('w:vAlign')
                va.set(qn('w:val'), 'center')
                tcPr0.append(va)
                tc0.append(make_cell_paragraph(scenario, align="center"))
            else:
                vm.set(qn('w:val'), 'continue')
                tcPr0.append(vm)
                # Empty paragraph
                tc0.append(OxmlElement('w:p'))
            set_col_width(tc0, col_widths[0])
            set_cell_padding(tc0)
            tr.append(tc0)

            # Columns 1-3
            for ci, val in enumerate([method, mse, sd]):
                tc = OxmlElement('w:tc')
                set_col_width(tc, col_widths[ci])
                align = "left" if ci == 0 else "center"
                tc.append(make_cell_paragraph(val, align=align))
                set_cell_padding(tc)
                tr.append(tc)

            tbl.append(tr)

    return tbl


# ══════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════

def main():
    # Step 1: Clone the complete file
    print(f"Step 1: Cloning {SOURCE} → {OUTPUT}")
    shutil.copyfile(SOURCE, OUTPUT)
    original_size = os.path.getsize(OUTPUT)
    print(f"  Cloned: {original_size} bytes")

    # Step 2: Open and modify in-place
    print("\nStep 2: Opening cloned file for in-place surgery...")
    doc = Document(OUTPUT)

    # ── 2a: Global formatting: TNR 10pt + exactly 12pt spacing ──
    print("  2a: Applying global formatting (paragraphs)...")
    for i, para in enumerate(doc.paragraphs):
        format_paragraph(para)
    print(f"    Formatted {len(doc.paragraphs)} paragraphs")

    # ── 2b: Find and replace Table 1, Table 2 ──
    print("\n  2b: Processing tables...")
    new_tbl1_xml = build_table1_xml()
    new_tbl2_xml = build_table2_xml()

    table1_replaced = False
    table2_replaced = False

    for ti, table in enumerate(doc.tables):
        # Identify table by checking first row text
        first_row_text = ""
        for cell in table.rows[0].cells:
            first_row_text += cell.text.strip().lower() + " "

        print(f"    Table {ti}: first row = {first_row_text[:80]}...")

        if "scenario" in first_row_text and "purpose" in first_row_text:
            # Table 1: Simulation scenarios
            print(f"      → Matched as Table 1 (Simulation scenarios)")
            tbl_elem = table._tbl

            # Remove all existing rows
            for tr in list(tbl_elem.findall(qn('w:tr'))):
                tbl_elem.remove(tr)

            # Insert new rows from new_tbl1_xml
            for tr in new_tbl1_xml.findall(qn('w:tr')):
                tbl_elem.append(tr)

            # Apply three-line borders
            apply_three_line_borders(table)
            table1_replaced = True
            print(f"      → Replaced with three-line table + OMML formulas")

        elif "scenario" in first_row_text and "method" in first_row_text and "mse" in first_row_text:
            # Table 2: Test MSE
            print(f"      → Matched as Table 2 (Test MSE)")
            tbl_elem = table._tbl

            for tr in list(tbl_elem.findall(qn('w:tr'))):
                tbl_elem.remove(tr)

            for tr in new_tbl2_xml.findall(qn('w:tr')):
                tbl_elem.append(tr)

            apply_three_line_borders(table)
            table2_replaced = True
            print(f"      → Replaced with three-line table + vertical merge")

    if not table1_replaced:
        print("    WARNING: Table 1 not found!")
    if not table2_replaced:
        print("    WARNING: Table 2 not found!")

    # ── 2c: Format remaining tables (3, 4, 5) ──
    print("\n  2c: Formatting remaining tables...")
    for ti, table in enumerate(doc.tables):
        # Skip already-processed tables
        first_row_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if ("scenario" in first_row_text and "purpose" in first_row_text) or \
           ("scenario" in first_row_text and "method" in first_row_text and "mse" in first_row_text):
            continue

        print(f"    Table {ti}: three-line formatting")
        apply_three_line_borders(table)
        # Format all cells
        for row in table.rows:
            for cell in row.cells:
                set_cell_padding(cell)
                for para in cell.paragraphs:
                    format_paragraph(para)

    # ── 2d: Ensure margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Step 3: Save ──
    print(f"\nStep 3: Saving to {OUTPUT}...")
    doc.save(OUTPUT)
    final_size = os.path.getsize(OUTPUT)
    print(f"  Saved: {final_size} bytes")

    # ── Step 4: Verify ──
    print("\nStep 4: Verification...")
    verify_doc = Document(OUTPUT)

    total_words = 0
    for p in verify_doc.paragraphs:
        t = p.text.strip()
        if t:
            total_words += len(t.split())

    print(f"  Paragraphs: {len(verify_doc.paragraphs)}")
    print(f"  Tables: {len(verify_doc.tables)}")
    print(f"  Word count: ~{total_words}")

    # Check all sections present
    section_titles = ["Introduction", "Preliminaries", "Related Work",
                      "Methodology", "Results", "Real-Data",
                      "Discussion", "Conclusion", "References"]
    found = 0
    for title in section_titles:
        for p in verify_doc.paragraphs:
            if title.lower() in p.text.lower():
                found += 1
                break
    print(f"  Sections found: {found}/{len(section_titles)}")

    # Verify key paragraph not truncated
    for p in verify_doc.paragraphs:
        if p.text.strip().startswith("However, these guarantees"):
            full = p.text.strip()
            print(f"  'However' paragraph length: {len(full)} chars")
            if len(full) > 200:
                print(f"  STATUS: COMPLETE (starts with '{full[:50]}...')")
            else:
                print(f"  WARNING: Paragraph seems short: {full}")
            break

    # Check line spacing
    exact_count = 0
    for p in verify_doc.paragraphs:
        lsr = p.paragraph_format.line_spacing_rule
        if lsr == WD_LINE_SPACING.EXACTLY:
            exact_count += 1
    print(f"  Paragraphs with EXACTLY spacing: {exact_count}")

    # Font check
    non_tnr = 0
    for p in verify_doc.paragraphs:
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rF = rPr.find(qn('w:rFonts'))
                if rF is not None:
                    af = rF.get(qn('w:ascii'))
                    if af and af != FONT_NAME:
                        non_tnr += 1
    print(f"  Non-TNR runs: {non_tnr}")

    # Save size report
    print(f"\n  File: {OUTPUT}")
    print(f"  Size: {final_size} bytes ({final_size/1024:.1f} KB)")

    print("\nDone!")


if __name__ == "__main__":
    main()
