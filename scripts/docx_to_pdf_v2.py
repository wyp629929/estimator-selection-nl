"""
Convert paper_final_perfect.docx to PDF via styled HTML + weasyprint.
Properly handles tables and document structure.
"""

import os
from lxml import etree
from docx import Document
from weasyprint import HTML

DOCX = "/Users/wangyaoping/Desktop/paper_final_perfect.docx"
PDF_OUT = "/Users/wangyaoping/Desktop/paper_final_perfect.pdf"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document(DOCX)

# Build body element order: interleave paragraphs and tables
# python-docx stores tables separately from paragraphs, so we need
# to walk the XML body to get the correct order

body = doc.element.body
html_parts = []
html_parts.append("""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
@page { size: A4; margin: 1in; }
body {
  font-family: 'Times New Roman', Times, serif;
  font-size: 10pt;
  line-height: 1.2;
  text-align: justify;
  orphans: 3; widows: 3;
}
.title { font-size: 14pt; font-weight: bold; text-align: center; margin-bottom: 12pt; }
.author-info { text-align: center; font-size: 9pt; margin-bottom: 4pt; }
.abstract-label { font-weight: bold; font-size: 10pt; margin-top: 12pt; }
.abstract-text { font-size: 10pt; text-align: justify; margin-bottom: 6pt; }
.keywords { font-size: 10pt; margin-bottom: 12pt; }
.section-heading { font-weight: bold; font-size: 11pt; margin-top: 14pt; margin-bottom: 6pt; }
.subsection-heading { font-weight: bold; font-size: 10pt; margin-top: 10pt; margin-bottom: 4pt; }
.body-text { font-size: 10pt; margin: 0 0 3pt 0; text-align: justify; }
.table-label { font-weight: bold; font-size: 9pt; text-align: center; margin-top: 10pt; margin-bottom: 3pt; }
table.doc-table { border-collapse: collapse; width: 100%; margin: 4pt auto; font-size: 9pt; page-break-inside: avoid; }
table.doc-table th {
  border-top: 2px solid black;
  border-bottom: 0.75pt solid black;
  padding: 3pt 4pt;
  text-align: center;
  font-weight: bold;
}
table.doc-table td {
  padding: 2pt 4pt;
  text-align: center;
  border: none;
}
table.doc-table {
  border-top: 2px solid black;
  border-bottom: 2px solid black;
  margin-bottom: 6pt;
}
.ref-entry { font-size: 9pt; margin: 0 0 2pt 0; text-align: justify; padding-left: 24pt; text-indent: -24pt; }
.scenario-cell { text-align: left !important; }
</style>
</head><body>
""")

def esc(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

def get_cell_text(cell):
    """Extract text from a table cell."""
    return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())

# Walk body children to get correct order
table_children = [c for c in body.iterchildren() if c.tag == f'{{{W_NS}}}tbl']
table_index = 0

is_first_para = True
for child in body.iterchildren():
    tag = child.tag
    if '}' in tag:
        local = tag.split('}')[1]
    else:
        local = tag

    if local == 'p':
        # Paragraph - get text
        texts = [t.text or '' for t in child.iter(f'{{{W_NS}}}t')]
        text = ''.join(texts).strip()
        if not text:
            continue

        # Check if paragraph contains math (oMath elements)
        has_math = child.findall('{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')

        if is_first_para:
            # Title
            html_parts.append(f'<div class="title">{esc(text)}</div>\n')
            is_first_para = False
        elif text in ["Abstract"]:
            html_parts.append(f'<div class="abstract-label">Abstract</div>\n')
        elif text.startswith("Keywords:"):
            html_parts.append(f'<p class="keywords"><b>Keywords:</b> {esc(text[9:])}</p>\n')
        elif text.startswith("Contributions."):
            html_parts.append(f'<p class="body-text"><b>Contributions.</b> {esc(text[14:])}</p>\n')
        elif text.startswith("Table ") and len(text) < 60:
            html_parts.append(f'<p class="table-label">{esc(text)}</p>\n')
        elif text.startswith("Figure ") and len(text) < 60:
            html_parts.append(f'<p class="table-label">{esc(text)}</p>\n')
        elif any(text.startswith(f"{i}.") for i in range(0, 10)):
            if '.' in text[2:]:  # has subsection number like 2.1
                html_parts.append(f'<div class="subsection-heading">{esc(text)}</div>\n')
            else:
                html_parts.append(f'<div class="section-heading">{esc(text)}</div>\n')
        elif any(text == s for s in ["Introduction", "Preliminaries", "Related Work",
            "Experimental Methodology", "Simulation Results", "Real-Data Applications",
            "Discussion", "Conclusion", "References"]):
            html_parts.append(f'<div class="section-heading">{esc(text)}</div>\n')
        elif any(text.startswith(s) for s in ["Parameter Estimation in Classical Statistics",
            "Tree-Based Methods for Estimation", "Deep Neural Networks",
            "Statistical Properties of ML Estimators", "ML for Parameter Estimation",
            "Uncertainty Quantification for ML Methods",
            "Simulation Design", "Methods Compared", "Evaluation Protocol",
            "Linear Scenario", "Semiparametric Scenario", "Nonlinear Scenario",
            "High-Dimensional Sparse Scenario",
            "Financial Risk Modeling", "Biomedical Outcome Prediction",
            "Regimes Where ML Outperforms", "Computational Trade-offs", "Limitations"]):
            html_parts.append(f'<div class="subsection-heading">{esc(text)}</div>\n')
        elif text.startswith("[") and "]" in text[:5]:
            html_parts.append(f'<p class="ref-entry">{esc(text)}</p>\n')
        elif any(k in text for k in ["WANG", "School", "@", "United", "Rutgers"]):
            html_parts.append(f'<div class="author-info">{esc(text)}</div>\n')
        else:
            # Render math formulas as italic text approximations
            if has_math:
                # Try to render math as best we can
                html_parts.append(f'<p class="body-text">{esc(text)}</p>\n')
            else:
                html_parts.append(f'<p class="body-text">{esc(text)}</p>\n')

    elif local == 'tbl':
        # Table - render as HTML table
        if table_index >= len(doc.tables):
            table_index += 1
            continue

        table = doc.tables[table_index]
        rows_data = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(get_cell_text(cell))
            rows_data.append(cells)

        if not rows_data:
            table_index += 1
            continue

        is_scenario_table = "scenario" in rows_data[0][0].lower() if rows_data else False

        cls = "doc-table"
        html_parts.append(f'<table class="{cls}">\n')

        # Header
        html_parts.append('<tr>\n')
        for h in rows_data[0]:
            html_parts.append(f'<th>{esc(h)}</th>\n')
        html_parts.append('</tr>\n')

        # Data
        for row in rows_data[1:]:
            html_parts.append('<tr>\n')
            for ci, cell in enumerate(row):
                cls_attr = ' class="scenario-cell"' if ci == 0 or ci == 1 else ''
                html_parts.append(f'<td{cls_attr}>{esc(cell)}</td>\n')
            html_parts.append('</tr>\n')

        html_parts.append('</table>\n')
        table_index += 1

html_parts.append("</body></html>")

# Write HTML and convert
import tempfile
with tempfile.NamedTemporaryFile(suffix='.html', mode='w', delete=False, encoding='utf-8') as f:
    f.write(''.join(html_parts))
    html_path = f.name

print(f"HTML temp: {html_path}")
HTML(html_path).write_pdf(PDF_OUT)
os.unlink(html_path)

size = os.path.getsize(PDF_OUT)
# Count pages
try:
    from pypdf import PdfReader
    reader = PdfReader(PDF_OUT)
    pages = len(reader.pages)
except:
    pages = "?"
print(f"PDF saved: {PDF_OUT} ({size} bytes, {pages} pages)")
