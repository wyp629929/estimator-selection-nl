"""
Convert paper_final_perfect.docx to PDF using fpdf2.
"""

from docx import Document
from fpdf import FPDF
import os

DOCX = "/Users/wangyaoping/Desktop/paper_final_perfect.docx"
OUTPUT = "/Users/wangyaoping/Desktop/paper_final_perfect.pdf"

doc = Document(DOCX)

pdf = FPDF(orientation='P', unit='mm', format='A4')
pdf.add_page()
pdf.set_auto_page_break(auto=True, margin=25)

# Margins: 1 inch = 25.4 mm
MARGIN = 25.4
pdf.set_margins(MARGIN, MARGIN, MARGIN)

# Font settings
FONT = "Times"
LINE_H = 5.5  # ~12pt line spacing in mm


def write_body(text, bold=False, size=10, align='J', space_after=0.5):
    """Write body text with word wrap."""
    if not text.strip():
        pdf.ln(LINE_H)
        return
    pdf.set_font(FONT, 'B' if bold else '', size)
    pdf.multi_cell(w=0, h=LINE_H, text=text.strip(), align=align)
    if space_after > 0:
        pdf.ln(space_after)


def write_heading(text, level=1):
    """Write a heading."""
    if level == 0:  # Title
        pdf.set_font(FONT, 'B', 14)
        pdf.multi_cell(w=0, h=7, text=text.strip(), align='C')
        pdf.ln(2)
    elif level == 1:  # Section (1., 2., etc.)
        pdf.ln(1)
        pdf.set_font(FONT, 'B', 11)
        pdf.multi_cell(w=0, h=LINE_H + 1, text=text.strip(), align='J')
    elif level == 2:  # Subsection (1.1., 2.1., etc.)
        pdf.set_font(FONT, 'B', 10)
        pdf.multi_cell(w=0, h=LINE_H, text=text.strip(), align='J')
    else:
        write_body(text, bold=True)


def write_table(table_data, headers):
    """Write a simple table."""
    col_w = (210 - 2 * MARGIN) / len(headers)

    # Header
    pdf.set_font(FONT, 'B', 9)
    for h, cw in zip(headers, [col_w] * len(headers)):
        pdf.cell(cw, LINE_H + 1, h, border=1, align='C')
    pdf.ln()

    # Data
    pdf.set_font(FONT, '', 9)
    for row in table_data:
        for cell, cw in zip(row, [col_w] * len(row)):
            pdf.cell(cw, LINE_H + 1, str(cell)[:20], border=1, align='C')
        pdf.ln()


# ── Parse sections ──
section_keywords = {
    "Abstract": 0,
    "Introduction": 1,
    "Preliminaries": 1,
    "Related Work": 1,
    "Experimental Methodology": 1,
    "Simulation Results": 1,
    "Real-Data Applications": 1,
    "Discussion": 1,
    "Conclusion": 1,
    "References": 1,
}
subsection_prefixes = [
    "Parameter Estimation", "Tree-Based Methods", "Deep Neural Networks",
    "Statistical Properties", "ML for Parameter", "Uncertainty Quantification",
    "Simulation Design", "Methods Compared", "Evaluation Protocol",
    "Linear Scenario", "Semiparametric Scenario", "Nonlinear Scenario",
    "High-Dimensional Sparse Scenario",
    "Financial Risk", "Biomedical Outcome",
    "Regimes Where", "Computational Trade", "Limitations",
]

def detect_level(text):
    for k in section_keywords:
        if text.startswith(k) or text == k:
            return section_keywords[k]
    for p in subsection_prefixes:
        if text.startswith(p):
            return 2
    return None

# ── Title page ──
para0 = doc.paragraphs[0].text.strip()
write_heading(para0, 0)

for p in doc.paragraphs[1:5]:
    t = p.text.strip()
    if t:
        if any(k in t for k in ["WANG", "School", "@", "United"]):
            write_body(t, size=9, align='C', space_after=0.3)
        else:
            write_body(t, size=10, align='C', space_after=0.5)

pdf.ln(2)

# ── Body ──
table_idx = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue

    level = detect_level(t)

    if level is not None:
        write_heading(t, level)
    elif t.startswith("[") and "]" in t[:5]:
        # Reference
        write_body(t, size=9, space_after=0.3)
    else:
        write_body(t, size=10, space_after=0.3)

# ── Tables ──
# Add tables from docx as text representations
for ti, table in enumerate(doc.tables):
    if pdf.get_y() > 250:
        pdf.add_page()
    pdf.ln(1)

    # Extract table data
    rows_data = []
    for row in table.rows:
        rows_data.append([cell.text.strip()[:30] for cell in row.cells])

    if rows_data:
        headers = rows_data[0]
        data = rows_data[1:]
        write_table(data, headers)

    pdf.ln(2)

pdf.output(OUTPUT)
print(f"PDF saved: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT)} bytes")
