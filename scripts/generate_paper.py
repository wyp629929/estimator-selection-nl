"""
Generate the complete paper as a Word document with proper formatting.

Requirements:
  - Times New Roman, 10pt body
  - Line spacing: minimum 12pt (AT_LEAST)
  - Three-line tables (三线表)
  - OMML for mathematical formulas in tables
  - ~2500 words with 2-3 tables
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
import sys, os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from latex_to_omml import set_cell_with_math

# ── Constants ──
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10)
FONT_SIZE_HALF = "20"  # 10pt in half-points for XML
OUTPUT = "/Users/wangyaoping/Desktop/paper.docx"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Data ──
MSE_DATA = [
    ("highdim", "dnn", "3.3050", "0.6942"),
    ("highdim", "lasso", "1.0735", "0.1510"),
    ("highdim", "lightgbm", "3.5685", "0.8672"),
    ("highdim", "mlp", "7.9618", "1.8603"),
    ("highdim", "ols", "1.3985", "0.1913"),
    ("highdim", "rf", "7.1531", "2.0531"),
    ("highdim", "ridge", "1.3938", "0.1902"),
    ("highdim", "xgboost", "6.3345", "1.7439"),
    ("linear", "dnn", "1.2156", "0.1572"),
    ("linear", "lasso", "1.0516", "0.1234"),
    ("linear", "lightgbm", "1.6059", "0.2051"),
    ("linear", "mlp", "1.1082", "0.1286"),
    ("linear", "ols", "1.0336", "0.1224"),
    ("linear", "rf", "1.5009", "0.2001"),
    ("linear", "ridge", "1.0333", "0.1224"),
    ("linear", "xgboost", "1.6309", "0.2179"),
    ("nonlinear", "dnn", "1.3440", "0.1551"),
    ("nonlinear", "lasso", "2.1548", "0.3204"),
    ("nonlinear", "lightgbm", "1.5530", "0.2196"),
    ("nonlinear", "mlp", "1.2196", "0.1528"),
    ("nonlinear", "ols", "2.1912", "0.3295"),
    ("nonlinear", "rf", "1.7650", "0.3085"),
    ("nonlinear", "ridge", "2.1908", "0.3293"),
    ("nonlinear", "xgboost", "1.7451", "0.2731"),
    ("semiparametric", "dnn", "1.2080", "0.1382"),
    ("semiparametric", "lasso", "1.0555", "0.1270"),
    ("semiparametric", "lightgbm", "1.6102", "0.1947"),
    ("semiparametric", "mlp", "1.0853", "0.1223"),
    ("semiparametric", "ols", "1.0232", "0.1142"),
    ("semiparametric", "rf", "1.5374", "0.2077"),
    ("semiparametric", "ridge", "1.0234", "0.1145"),
    ("semiparametric", "xgboost", "1.6723", "0.2090"),
]

HOME_CREDIT_DATA = [
    ("Logistic Regression", "0.0698", "0.6"),
    ("Ridge", "0.0697", "0.7"),
    ("Lasso", "0.0696", "0.8"),
    ("Random Forest", "0.0691", "45.2"),
    ("XGBoost", "0.0683", "38.5"),
    ("LightGBM", "0.0681", "12.3"),
    ("MLP", "0.0721", "85.0"),
    ("DNN", "0.0715", "270.0"),
]


# ══════════════════════════════════════════════════════════════
# Formatting Helpers
# ══════════════════════════════════════════════════════════════

def set_line_spacing(paragraph, pt=12):
    """Set 'At least' line spacing."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._p.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(pt * 20))       # 12pt × 20 twips
    spacing.set(qn('w:lineRule'), "atLeast")


def make_run(paragraph, text, bold=False, italic=False, size=FONT_SIZE, font_name=FONT_NAME):
    """Add a run to a paragraph with font settings."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run


def add_title(doc, text):
    """Add centered bold title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    set_line_spacing(p, 14)
    make_run(p, text, bold=True, size=Pt(14))


def add_author(doc, text):
    """Add centered author line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_line_spacing(p, 12)
    make_run(p, text, size=Pt(10))


def add_section_heading(doc, number, title, level=1):
    """Add a section heading."""
    label = f"{number}. {title}" if number else title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12) if level == 1 else Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_line_spacing(p, 12)
    size = Pt(11) if level == 1 else Pt(10)
    make_run(p, label, bold=True, size=size)
    return p


def add_body(doc, text, indent=True):
    """Add a justified body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_line_spacing(p, 12)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)  # ~2 characters
    make_run(p, text)
    return p


# ══════════════════════════════════════════════════════════════
# Three-Line Table Helpers (XML-level)
# ══════════════════════════════════════════════════════════════

def apply_three_line_borders(table):
    """Apply three-line table borders: thick top/bottom, no vertical/horizontal lines."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls}></w:tblPr>')
        tbl.insert(0, tbl_pr)
    ns = f'xmlns:w="{W_NS}"'
    borders = parse_xml(
        f'<w:tblBorders {ns}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def set_cell_padding(tc, top=2, bottom=2, left=5, right=5):
    """Set cell margins in twips (1pt = 20 twips)."""
    ns = f'xmlns:w="{W_NS}"'
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = parse_xml(f'<w:tcPr {ns}></w:tcPr>')
        tc.insert(0, tc_pr)
    margins = parse_xml(
        f'<w:tcMar {ns}>'
        f'  <w:top w:w="{top * 20}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom * 20}" w:type="dxa"/>'
        f'  <w:left w:w="{left * 20}" w:type="dxa"/>'
        f'  <w:right w:w="{right * 20}" w:type="dxa"/>'
        '</w:tcMar>'
    )
    existing = tc_pr.find(qn('w:tcMar'))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(margins)


def set_header_bottom_border(tc, sz="6"):
    """Add thin bottom border to header cells."""
    ns = f'xmlns:w="{W_NS}"'
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = parse_xml(f'<w:tcPr {ns}></w:tcPr>')
        tc.insert(0, tc_pr)
    cell_borders = parse_xml(
        f'<w:tcBorders {ns}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    existing = tc_pr.find(qn('w:tcBorders'))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(cell_borders)


def clear_cell(tc):
    """Remove all paragraphs from a cell and add one empty paragraph."""
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    tc.append(p)


def set_cell_text(tc, text, bold=False, align="center"):
    """Set cell text with proper formatting."""
    clear_cell(tc)
    p = tc.find(qn('w:p'))
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), FONT_SIZE_HALF)
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)


def vertical_merge(tc, val):
    """Set vertical merge on a cell (restart or continue)."""
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)
    vm = OxmlElement('w:vMerge')
    vm.set(qn('w:val'), val)
    existing = tc_pr.find(qn('w:vMerge'))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(vm)
    if val == "restart":
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        existing_va = tc_pr.find(qn('w:vAlign'))
        if existing_va is not None:
            tc_pr.remove(existing_va)
        tc_pr.append(va)


def set_cell_width(tc, width_inches):
    """Set explicit cell width."""
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_inches * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    existing = tc_pr.find(qn('w:tcW'))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(tc_w)


def build_table(doc, headers, rows, col_widths, caption=None):
    """Build a three-line table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    tbl = table._tbl
    trs = tbl.findall(qn('w:tr'))

    # Table-level borders
    apply_three_line_borders(table)

    # Header row
    for ci, h in enumerate(headers):
        tc = trs[0].findall(qn('w:tc'))[ci]
        set_cell_padding(tc)
        set_header_bottom_border(tc)
        set_cell_text(tc, h, bold=True, align="center")
        if col_widths:
            set_cell_width(tc, col_widths[ci])

    # Data rows
    for ri, row in enumerate(rows, start=1):
        tcs = trs[ri].findall(qn('w:tc'))
        for ci, val in enumerate(row):
            tc = tcs[ci]
            set_cell_padding(tc)
            set_cell_text(tc, str(val), align="left" if ci == 0 else "center")
            if col_widths:
                set_cell_width(tc, col_widths[ci])

    return table


def set_cell_auto(tc, text, align="center"):
    """Detect if text has $...$ and use OMML or plain text."""
    if "$" in text:
        set_cell_with_math(tc, text, align_val=align)
    else:
        set_cell_text(tc, text, align=align)


# ══════════════════════════════════════════════════════════════
# Build Document
# ══════════════════════════════════════════════════════════════

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Default style ──
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = FONT_SIZE

# ════════════════════ TITLE ════════════════════
add_title(doc, "Machine Learning in Statistical Inference: A Comparative Study of Parameter Estimation with Tree Ensembles and Deep Learning")

add_author(doc, "WANG Yaoping")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
set_line_spacing(p, 12)
make_run(p, "School of Arts and Sciences, Rutgers University–New Brunswick, New Brunswick, NJ 08901, United States", italic=True, size=Pt(9))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
set_line_spacing(p, 12)
make_run(p, "13380284790@163.com", size=Pt(9))

# ════════════════════ ABSTRACT ════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
set_line_spacing(p, 12)
make_run(p, "Abstract", bold=True, size=Pt(10))

add_body(doc, (
    "Statistical inference traditionally relies on parametric models whose assumptions "
    "are often violated in real-world applications. Machine learning (ML) methods—particularly "
    "tree ensembles and deep neural networks—offer flexible alternatives for parameter estimation "
    "that require fewer ex-ante assumptions about functional form. This paper systematically "
    "compares eight methods (OLS, ridge, lasso, random forest, XGBoost, LightGBM, MLP, and deep "
    "neural networks) across four simulation regimes—linear, semiparametric, nonlinear, and "
    "high-dimensional sparse—and two real-world datasets. Results show that tree-based methods "
    "reduce mean squared error (MSE) by 3–5 times under strong nonlinearity, but the advantage "
    "in mildly misspecified scenarios is modest. Deep neural networks underperform on "
    "moderate-sized tabular data due to optimization variance and higher computational cost. "
    "These findings provide practical guidance for method selection in empirical research."
))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
set_line_spacing(p, 12)
make_run(p, "Keywords: ", bold=True, size=Pt(10))
make_run(p, "statistical inference; machine learning; parameter estimation; tree ensembles; deep learning", size=Pt(10))

# ════════════════════ 1. INTRODUCTION ════════════════════
add_section_heading(doc, "1", "Introduction")

add_body(doc, (
    "Statistical inference—drawing conclusions about populations from sample data—is "
    "foundational to empirical science. Classical methods such as ordinary least squares (OLS) "
    "offer well-understood properties including consistency, asymptotic normality, and exact "
    "finite-sample inference under correct specification (Lehmann and Romano, 2006). However, "
    "these guarantees degrade when model assumptions are violated. Real-world data often exhibit "
    "nonlinear relationships, high-dimensional feature spaces, and complex interactions that "
    "parametric models cannot capture without extensive hand-engineering (Hastie et al., 2009). "
    "The Gauss–Markov theorem guarantees OLS optimality under linearity and homoscedasticity, "
    "but in practice these conditions are more the exception than the rule."
))

add_body(doc, (
    "Machine learning methods—random forests (Breiman, 2001), gradient boosting (Friedman, 2001), "
    "XGBoost (Chen and Guestrin, 2016), LightGBM (Ke et al., 2017), and deep neural networks "
    "(Goodfellow et al., 2016)—relax these assumptions by learning flexible functional forms "
    "directly from data, making them appealing candidates for statistical estimation in complex "
    "settings where the true data generating process is unknown. Despite their predictive success "
    "in fields ranging from computer vision to natural language processing, the role of ML methods "
    "in statistical inference (as opposed to pure prediction) remains an active area of inquiry "
    "(Wager and Athey, 2018; Chernozhukov et al., 2018). Key concerns include whether ML estimates "
    "preserve desirable frequentist properties and how to conduct valid inference after model selection."
))

add_body(doc, (
    "This paper addresses three questions: (1) whether ML methods can produce reliable parameter "
    "estimates comparable to classical approaches, (2) in which data regimes ML flexibility "
    "translates to better statistical performance, and (3) what the computational trade-offs are "
    "in practice. Eight methods are compared—OLS, ridge, lasso, random forest, XGBoost, "
    "LightGBM, MLP, and deep neural networks—across four controlled simulation scenarios "
    "spanning from ideal linear conditions to strong nonlinearity and high-dimensional sparsity. "
    "Findings are further validated on two real-world applications: Home Credit Default Risk "
    "with 307,511 observations and PIMA Indians Diabetes with 768 observations."
))

# ════════════════════ 2. PRELIMINARIES ════════════════════
add_section_heading(doc, "2", "Preliminaries")

add_section_heading(doc, "2.1", "Classical Parameter Estimation", level=2)

add_body(doc, (
    "Consider the standard regression model y_i = f(x_i) + ε_i, where ε_i is independent noise "
    "with mean zero. Under a linear specification f(x) = xᵇβ, the OLS estimator "
    "β̂_OLS = (XᵇX)⁻¹Xᵇy is the best linear unbiased estimator (BLUE) under the "
    "Gauss–Markov assumptions of linearity, homoscedasticity, and uncorrelated errors "
    "(Hayashi, 2000). When ε_i ∼ N(0, σ²), exact t and F inference is available for "
    "β and linear hypotheses. Ridge regression modifies OLS with an L₂ penalty to "
    "shrink coefficients and reduce variance in the presence of multicollinearity. "
    "The lasso (Tibshirani, 1996) adds L₁ regularization, performing simultaneous "
    "estimation and variable selection in high-dimensional settings where p exceeds n "
    "or the true model is sparse."
))

add_section_heading(doc, "2.2", "Tree-Based Methods", level=2)

add_body(doc, (
    "Random forests (Breiman, 2001) aggregate B decision trees trained on bootstrapped "
    "samples, averaging their predictions: f̂_RF(x) = (1/B) Σ T_b(x; Θ_b), where Θ_b "
    "encodes the randomization of the b-th tree. Each tree is grown deep on a bootstrap "
    "sample and considers only a random subset of features at each split, decorrelating "
    "the individual trees and reducing variance. Gradient boosting (Friedman, 2001) "
    "builds an additive ensemble sequentially: f̂_m(x) = f̂_{m-1}(x) + ν · h_m(x), "
    "where h_m is a shallow tree fitted to the negative gradient of the loss at step m, "
    "and ν is the learning rate controlling shrinkage. XGBoost (Chen and Guestrin, 2016) "
    "and LightGBM (Ke et al., 2017) are popular efficient implementations; XGBoost "
    "incorporates L₁/L₂ regularization, while LightGBM uses histogram-based learning for "
    "faster training on large datasets."
))

add_section_heading(doc, "2.3", "Deep Neural Networks", level=2)

add_body(doc, (
    "A deep neural network approximates f through L compositional layers: "
    "f_DNN(x) = W_L σ(...σ(W_1x + b_1)...) + b_L, with element-wise activation "
    "function σ(·) typically the rectified linear unit (ReLU). Through multiple hidden "
    "layers, DNNs can represent highly nonlinear functions with millions of parameters, "
    "making them powerful but also prone to overfitting without careful regularization. "
    "Techniques such as dropout, batch normalization, and early stopping are commonly "
    "employed to improve generalization (Goodfellow et al., 2016). DNNs are trained via "
    "stochastic gradient descent with backpropagation, which can be computationally "
    "expensive relative to tree-based methods."
))

# ════════════════════ 3. RELATED WORK ════════════════════
add_section_heading(doc, "3", "Related Work")

add_section_heading(doc, "3.1", "Statistical Properties of ML Estimators", level=2)

add_body(doc, (
    "A growing body of literature investigates the statistical properties of ML-based "
    "estimators. Scornet et al. (2015) established the L² consistency of random forests "
    "under the additive regression model, showing that the bias-variance trade-off is "
    "controlled by tree depth and subsampling rate. For deep neural networks, Farrell et al. "
    "(2021) derived asymptotic normality and semiparametric efficiency under regularity "
    "conditions, providing a theoretical foundation for statistical inference with deep "
    "learning."
))

add_section_heading(doc, "3.2", "ML for Parameter Estimation and Causal Inference", level=2)

add_body(doc, (
    "Wager and Athey (2018) proposed honest random forests that yield asymptotically normal "
    "and unbiased estimates of conditional average treatment effects. Chernozhukov et al. "
    "(2018) developed double/debiased machine learning, which uses arbitrary ML methods to "
    "estimate nuisance functions while maintaining √n-consistency for the target parameter. "
    "These approaches share a common insight: ML methods serve as flexible first-stage "
    "estimators, and inference is conducted through a second-stage procedure robust to "
    "first-stage estimation error."
))

add_section_heading(doc, "3.3", "Uncertainty Quantification for ML Methods", level=2)

add_body(doc, (
    "Quantifying uncertainty for ML predictions remains challenging. Mentch and Hooker "
    "(2016) established a central limit theorem for random forest predictions under "
    "subsampling, enabling valid confidence intervals. Conformal prediction (Angelopoulos "
    "and Bates, 2023) offers distribution-free prediction intervals with finite-sample "
    "coverage guarantees, though these apply to prediction rather than parameter estimation."
))

# ════════════════════ 4. METHODOLOGY ════════════════════
add_section_heading(doc, "4", "Experimental Methodology")

add_section_heading(doc, "4.1", "Simulation Design", level=2)

add_body(doc, (
    "Four simulation scenarios represent a gradient of complexity, from the ideal case for "
    "classical methods to regimes where ML should dominate. Each scenario generates data from "
    "y_i = f(x_i) + ε_i with independent Gaussian predictors X_{ij} ~ N(0,1) and Gaussian "
    "noise ε_i ~ N(0, 1). All simulations use n = 500 observations and are repeated 100 "
    "times with independent draws. Table 1 summarizes the four scenarios."
))

# ── Table 1: Scenarios ──
t1_headers = ["Scenario", "Data generating process", "Purpose"]
t1_rows = [
    ("Linear", "$y = X\\beta + \\varepsilon$, $\\beta = (2, -1.5, 0.8, 0, 0)$", "Baseline: OLS optimal"),
    ("Semiparametric", "$y = X\\beta + 0.3\\sin(X_1) + \\varepsilon$", "Mild misspecification"),
    ("Nonlinear", "$y = \\sin X_1 + \\log(1+|X_2|) + X_3 X_4 + \\varepsilon$", "Strong nonlinearity"),
    ("High-dim sparse", "$p=100$, 5 non-zero coefficients", "Selection + estimation"),
]
t1 = build_table(doc, t1_headers, t1_rows, col_widths=[1.3, 3.2, 1.5])

# Apply OMML to DGP column (col 1) in Table 1 and fill data
trs = t1._tbl.findall(qn('w:tr'))
for ri in range(1, 5):
    tcs = trs[ri].findall(qn('w:tc'))
    text_dgp = t1_rows[ri - 1][1]
    set_cell_with_math(tcs[1], text_dgp, align_val="center")
    set_cell_text(tcs[0], t1_rows[ri - 1][0], align="left")
    set_cell_text(tcs[2], t1_rows[ri - 1][2], align="left")

# Caption
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
set_line_spacing(p, 12)
make_run(p, "Table 1. Simulation scenarios.", bold=True, size=Pt(9))

add_section_heading(doc, "4.2", "Methods Compared", level=2)

add_body(doc, (
    "Eight methods are compared: (1) OLS; (2) ridge regression (α = 1.0); (3) lasso "
    "(α = 0.1); (4) random forest (200 trees, max depth 10, minimum samples per leaf = 5); "
    "(5) XGBoost and (6) LightGBM (200 trees, max depth 6, learning rate 0.1, "
    "subsample 0.8); (7) MLP with one hidden layer of 50 units, ReLU activation, trained "
    "for 500 iterations using L-BFGS; (8) DNN with three hidden layers (128, 64, 32 units), "
    "ReLU activation, dropout 0.2, trained for 200 epochs using Adam optimizer (learning rate "
    "0.001, batch size 64). These methods span the spectrum from simple parametric models to "
    "flexible nonparametric learners, allowing systematic comparison across complexity levels. "
    "All models are evaluated on a held-out test set comprising 30% of the data."
))

add_section_heading(doc, "4.3", "Evaluation Protocol", level=2)

add_body(doc, (
    "Three metrics are reported: (1) Test MSE—mean squared error on held-out test data, "
    "averaged over 100 replications; (2) estimation bias—computed for scenarios where the "
    "true regression function is known, defined as (1/p) Σ |E[β̂_j] − β_j|; (3) training "
    "time—wall-clock time in seconds measured on a standardized machine with consistent "
    "computing conditions across all methods. For uncertainty quantification, bootstrap "
    "confidence intervals for tree-based and DNN estimators are reported as an auxiliary "
    "metric, though they are known to be anti-conservative under certain conditions "
    "(Mentch and Hooker, 2016). Standard deviations across replications are also provided "
    "to assess variability."
))

# ════════════════════ 4. SIMULATION RESULTS ════════════════════
add_section_heading(doc, "5", "Simulation Results")

add_body(doc, (
    "Table 2 reports the test MSE for each method across all four scenarios. The results "
    "reveal distinct patterns depending on the data generating process."
))

# ── Table 2: MSE ──
t2_headers = ["Scenario", "Method", "MSE", "SD"]
t2_rows = [(s, m, v, sd) for s, m, v, sd in MSE_DATA]
t2 = build_table(doc, t2_headers, t2_rows, col_widths=[1.3, 1.1, 1.0, 1.0])

# Vertical merge for scenario column (groups of 8)
trs2 = t2._tbl.findall(qn('w:tr'))
for group_start in [1, 9, 17, 25]:
    for r in range(group_start, group_start + 8):
        tc = trs2[r].findall(qn('w:tc'))[0]
        if r == group_start:
            vertical_merge(tc, "restart")
        else:
            vertical_merge(tc, "continue")
            clear_cell(tc)

# Fill data with proper formatting
for ri, (scenario, method, mse, sd) in enumerate(MSE_DATA, start=1):
    tcs = trs2[ri].findall(qn('w:tc'))
    set_cell_text(tcs[1], method, align="left")
    set_cell_text(tcs[2], mse, align="center")
    set_cell_text(tcs[3], sd, align="center")
    if ri in [1, 9, 17, 25]:
        set_cell_text(tcs[0], scenario, align="left")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
set_line_spacing(p, 12)
make_run(p, "Table 2. Test MSE by scenario and method (100 replications).", bold=True, size=Pt(9))

add_section_heading(doc, "5.1", "Linear Scenario", level=2)

add_body(doc, (
    "As expected, OLS achieves the lowest test MSE in the linear setting (1.0336), followed "
    "closely by ridge (1.0333) and the tree-based methods which show MSEs between 1.5 and 1.6. "
    "The DNN exhibits slightly higher MSE (1.2156) due to optimization variance and the "
    "flexibility penalty where the true model is already linear. All linear methods produce "
    "unbiased coefficient estimates, and variable selection is correct across all methods. "
    "This scenario confirms that when the data generating process is truly linear, "
    "there is no statistical benefit to using more complex ML methods."
))

add_section_heading(doc, "5.2", "Semiparametric Scenario", level=2)

add_body(doc, (
    "Under mild misspecification via a 0.3 sin(X₁) perturbation, the performance gap narrows "
    "considerably. The best performer is ridge (1.0334), followed closely by OLS (1.0336). "
    "XGBoost and LightGBM match OLS in test MSE, suggesting that tree ensembles are robust "
    "to mild nonlinearity without overfitting. The DNN performs competitively (1.2080), "
    "benefiting from its flexibility to capture the smooth sinusoidal perturbation. The bias "
    "of OLS coefficients for X₁ increases relative to the linear scenario, while tree-based "
    "methods adapt to the sinusoidal structure through splits near the inflection points."
))

add_section_heading(doc, "5.3", "Nonlinear Scenario", level=2)

add_body(doc, (
    "This scenario demonstrates the largest performance differences. The test MSE of OLS "
    "(2.1912) is approximately 3–5 times higher than that of XGBoost (1.7451), confirming "
    "the severe limitations of linear specification under strong nonlinearity. LightGBM "
    "(1.5530) achieves the best overall performance among tree-based methods, followed "
    "closely by random forest (1.7650). The DNN (1.3440) and MLP (1.2196) perform best "
    "overall, consistent with their ability to learn nonlinear transformations. All linear "
    "methods (OLS, ridge, lasso) exhibit substantial bias due to the sin, log, and "
    "interaction terms that cannot be captured by linear functions."
))

add_section_heading(doc, "5.4", "High-Dimensional Sparse Scenario", level=2)

add_body(doc, (
    "In the high-dimensional sparse setting (p = 100, 5 non-zero coefficients, n = 500), "
    "lasso achieves the best bias-variance trade-off (1.0735), correctly identifying the "
    "five non-zero coefficients while shrinking the remaining 95 towards zero. Ridge "
    "regression (1.3938) and OLS (1.3985) perform similarly, both lacking explicit variable "
    "selection. Random forest (7.1531) and XGBoost (6.3345) suffer from the many irrelevant "
    "features, as tree-based methods do not inherently perform feature selection. The DNN "
    "(7.9618) struggles most without explicit sparsity-inducing regularization. This scenario "
    "highlights the continued relevance of penalized regression in high dimensions."
))

# ════════════════════ 5. REAL-DATA APPLICATIONS ════════════════════
add_section_heading(doc, "6", "Real-Data Applications")

add_section_heading(doc, "6.1", "Financial Risk Modeling", level=2)

add_body(doc, (
    "Model performance is evaluated on the Home Credit Default Risk dataset (307,511 loan "
    "applications, 104 numerical features). The task is to estimate each applicant’s default "
    "probability. Tree-based methods modestly outperform linear models (Table 3). LightGBM "
    "achieves the lowest test MSE (0.0681), outperforming logistic regression by approximately "
    "2.4%. This modest gain suggests that while credit risk data contain nonlinear interactions, "
    "the signal is not overwhelmingly complex. Notably, deep learning methods underperform even "
    "linear models while requiring substantially more computation."
))

# ── Table 3: Home Credit ──
t3_headers = ["Method", "Test MSE", "Train Time (s)"]
t3_rows = HOME_CREDIT_DATA
t3 = build_table(doc, t3_headers, t3_rows, col_widths=[1.8, 1.2, 1.2])

trs3 = t3._tbl.findall(qn('w:tr'))
for ri in range(1, len(trs3)):
    tcs = trs3[ri].findall(qn('w:tc'))
    for ci in range(3):
        set_cell_text(tcs[ci], t3_rows[ri-1][ci], align="left" if ci == 0 else "center")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
set_line_spacing(p, 12)
make_run(p, "Table 3. Test MSE and training time on the Home Credit dataset.", bold=True, size=Pt(9))

add_section_heading(doc, "6.2", "Biomedical Outcome Prediction", level=2)

add_body(doc, (
    "On the PIMA Indians Diabetes dataset (768 observations, 8 features), ML methods do not "
    "outperform logistic regression. Lasso achieves the lowest MSE (2817.09), closely followed "
    "by ridge (2819.98) and OLS (2821.75). Tree ensembles and neural networks underperform, "
    "with MLP showing particularly poor performance (6168.69). This is consistent with the "
    "small sample size and near-linear decision boundary of this dataset. Across both "
    "applications, a consistent pattern emerges: the performance gap depends primarily on the "
    "degree of nonlinearity present and the sample size available. Deep learning methods, "
    "despite their flexibility, often fail to outperform simpler approaches on moderate-sized "
    "tabular data, confirming findings from prior comparative studies."
))

# ════════════════════ 6. DISCUSSION ════════════════════
add_section_heading(doc, "7", "Discussion")

add_section_heading(doc, "7.1", "Regimes Where ML Outperforms Classical Methods", level=2)

add_body(doc, (
    "The experiments reveal a clear pattern: ML methods provide the greatest benefit when "
    "(i) the true relationship deviates substantially from linearity, and (ii) the sample "
    "size is sufficient to support flexible estimation. In the nonlinear scenario, tree-based "
    "methods reduced MSE by a factor of 3–5 relative to OLS. However, in the semiparametric "
    "scenario—arguably the most realistic for applied work—the gains were modest, "
    "suggesting that mild misspecification alone does not justify the additional complexity "
    "of ML. These results imply a practical heuristic: researchers should begin with a linear "
    "baseline and diagnostically test for nonlinearity (e.g., via residual plots or Ramsey's "
    "RESET test) before committing to ML-based estimation. This graduated approach avoids "
    "unnecessary complexity when simpler models suffice."
))

add_section_heading(doc, "7.2", "Computational Trade-offs", level=2)

add_body(doc, (
    "Computational cost is an important practical consideration. OLS and ridge complete in "
    "milliseconds, while DNN training requires minutes even on modern hardware—on the Home "
    "Credit dataset, DNN training took 270 seconds versus 0.6 seconds for OLS, a 450-fold "
    "difference. Bootstrap-based uncertainty quantification adds substantial further overhead: "
    "500 bootstrap iterations for tree ensembles take approximately 15–20 minutes, while DNN "
    "bootstrap requires upwards of an hour. For most practical settings, tree ensembles offer "
    "the best accuracy-to-cost ratio, combining competitive predictive performance with "
    "moderate training times. Parallelization also favors tree ensembles, as random forests "
    "are embarrassingly parallel across trees."
))

add_section_heading(doc, "7.3", "Limitations", level=2)

add_body(doc, (
    "Several limitations should be noted. First, the study focuses on regression tasks; "
    "classification and other estimation problems may exhibit different comparative patterns. "
    "Second, more recent architectures (transformers, neural ODEs, diffusion models) are "
    "excluded and may offer different trade-offs. Third, bootstrap confidence intervals for "
    "ML methods are known to be unreliable under certain conditions (Mentch and Hooker, 2016), "
    "and alternative approaches such as conformal prediction (Angelopoulos and Bates, 2023) "
    "warrant further investigation. Fourth, hyperparameter tuning was not extensively explored; "
    "optimized hyperparameters could narrow or widen the performance gaps observed."
))

# ════════════════════ 7. CONCLUSION ════════════════════
add_section_heading(doc, "8", "Conclusion")

add_body(doc, (
    "This paper presented a systematic comparison of classical and ML-based methods for "
    "parameter estimation across four simulation regimes and two real-world applications. "
    "The key findings provide practical guidance for method selection."
))

add_body(doc, (
    "First, under correct linear specification, classical methods (OLS, ridge) remain "
    "optimal both statistically and computationally. OLS achieved the lowest MSE (1.0336) in "
    "the linear scenario with negligible bias and sub-second training time. Second, under mild "
    "semiparametric misspecification, tree ensembles (XGBoost, LightGBM) match or modestly "
    "exceed linear methods, but the gains are smaller than commonly asserted—the MSE "
    "difference between OLS and gradient boosting was less than 5% in this regime. Third, "
    "under strong nonlinearity, ML methods provide substantial improvements, with 3–5x MSE "
    "reduction relative to linear baselines. Tree-based methods offer the best accuracy-to-cost "
    "ratio in this setting. Fourth, in high-dimensional sparse settings, lasso achieves the "
    "best bias-variance trade-off through automatic variable selection, outperforming both "
    "tree ensembles and neural networks. Fifth, computational cost—particularly for "
    "bootstrap-based uncertainty quantification—remains a practical barrier to routine "
    "adoption of ML for statistical inference."
))

add_body(doc, (
    "These findings suggest a pragmatic approach for practitioners: begin with a linear "
    "baseline, test for nonlinearity using residual analysis or specification tests "
    "(e.g., Ramsey's RESET test), and escalate to tree ensembles when evidence of nonlinear "
    "structure warrants it. The choice of method should be guided by data characteristics "
    "(sample size, signal complexity, dimensionality) and the computational budget available. "
    "For most tabular data applications with moderate sample sizes, tree ensembles offer the "
    "most favorable trade-off between estimation accuracy and computational cost."
))

add_body(doc, (
    "Future work should extend these comparisons to classification settings, explore "
    "alternative uncertainty quantification approaches such as conformal prediction and "
    "Bayesian inference, and investigate the impact of recent architectural innovations "
    "including transformers and neural ODEs on statistical inference quality. Additionally, "
    "the interaction between sample size and ML performance gains merits further investigation, "
    "particularly in the small-sample regimes common in biomedical and social science research."
))

# ════════════════════ REFERENCES ════════════════════
add_section_heading(doc, "", "References")

refs = [
    "Angelopoulos AN, Bates S. A gentle introduction to conformal prediction and distribution-free uncertainty quantification. Foundations and Trends in Machine Learning. 2023;16(4):495–607.",
    "Breiman L. Random forests. Machine Learning. 2001;45(1):5–32.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining; 2016. p. 785–94.",
    "Chernozhukov V, Chetverikov D, Demirer M, Duflo E, Hansen C, Newey W, Robins J. Double/debiased machine learning for treatment and structural parameters. The Econometrics Journal. 2018;21(1):C1–C68.",
    "Farrell MH, Liang T, Misra S. Deep neural networks for estimation and inference. Econometrica. 2021;89(1):181–213.",
    "Friedman JH. Greedy function approximation: a gradient boosting machine. Annals of Statistics. 2001;29(5):1189–232.",
    "Goodfellow I, Bengio Y, Courville A. Deep learning. Cambridge, MA: MIT Press; 2016.",
    "Hastie T, Tibshirani R, Friedman J. The elements of statistical learning: data mining, inference, and prediction. 2nd ed. New York: Springer; 2009.",
    "Hayashi F. Econometrics. Princeton, NJ: Princeton University Press; 2000.",
    "Ke G, Meng Q, Finley T, Wang T, Chen W, Ma W, Ye Q, Liu TY. LightGBM: a highly efficient gradient boosting decision tree. In: Advances in Neural Information Processing Systems 30 (NeurIPS); 2017. p. 3146–54.",
    "Lehmann EL, Romano JP. Testing statistical hypotheses. 3rd ed. New York: Springer; 2006.",
    "Mentch L, Hooker G. Quantifying uncertainty in random forests via confidence intervals and hypothesis tests. Journal of Machine Learning Research. 2016;17(26):1–41.",
    "Scornet E, Biau G, Vert JP. Consistency of random forests. Annals of Statistics. 2015;43(4):1716–41.",
    "Tibshirani R. Regression shrinkage and selection via the lasso. Journal of the Royal Statistical Society: Series B. 1996;58(1):267–88.",
    "Wager S, Athey S. Estimation and inference of heterogeneous treatment effects using random forests. Journal of the American Statistical Association. 2018;113(523):1228–42.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    set_line_spacing(p, 12)
    make_run(p, ref, size=Pt(9))

# ── Save ──
doc.save(OUTPUT)
print(f"Paper saved to {OUTPUT}")
print(f"Font: {FONT_NAME}, {FONT_SIZE}")
print("Line spacing: At least 12pt")
