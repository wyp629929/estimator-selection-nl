"""
Convert paper_final_perfect.docx to PDF via styled HTML + weasyprint.
"""

import os
from docx import Document
from weasyprint import HTML

DOCX = "/Users/wangyaoping/Desktop/paper_final_perfect.docx"
HTML_OUT = "/Users/wangyaoping/Desktop/paper_final_perfect.html"
PDF_OUT = "/Users/wangyaoping/Desktop/paper_final_perfect.pdf"

doc = Document(DOCX)

html_parts = []
html_parts.append("""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
@page {
  size: A4;
  margin: 1in;
}
body {
  font-family: 'Times New Roman', Times, serif;
  font-size: 10pt;
  line-height: 1.2;
  text-align: justify;
  orphans: 3;
  widows: 3;
}
.title {
  font-size: 14pt;
  font-weight: bold;
  text-align: center;
  margin-bottom: 12pt;
}
.author-info {
  text-align: center;
  font-size: 9pt;
  margin-bottom: 6pt;
}
.abstract-label {
  font-weight: bold;
  font-size: 10pt;
  margin-top: 12pt;
}
.abstract-text {
  font-size: 10pt;
  text-align: justify;
  margin-bottom: 6pt;
}
.keywords {
  font-size: 10pt;
  margin-bottom: 12pt;
}
.section-heading {
  font-weight: bold;
  font-size: 11pt;
  margin-top: 12pt;
  margin-bottom: 6pt;
}
.subsection-heading {
  font-weight: bold;
  font-size: 10pt;
  margin-top: 8pt;
  margin-bottom: 4pt;
}
.body-text {
  font-size: 10pt;
  text-indent: 0;
  margin: 0 0 4pt 0;
  text-align: justify;
}
.table-caption {
  font-weight: bold;
  font-size: 9pt;
  text-align: center;
  margin-top: 8pt;
  margin-bottom: 4pt;
}
table {
  border-collapse: collapse;
  width: 100%;
  margin: 6pt auto;
  font-size: 9pt;
}
table.scenario-table td, table.scenario-table th {
  border: 1px solid black;
  padding: 3pt 6pt;
  text-align: center;
}
table.scenario-table th {
  font-weight: bold;
  border-top: 2px solid black;
  border-bottom: 1px solid black;
}
table.scenario-table td {
  border: none;
}
table.scenario-table tr:first-child td {
  border-bottom: 0.5pt solid black;
}
table.scenario-table {
  border-top: 2px solid black;
  border-bottom: 2px solid black;
}
.ref-entry {
  font-size: 9pt;
  margin: 0 0 3pt 0;
  text-align: justify;
  padding-left: 24pt;
  text-indent: -24pt;
}
</style>
</head><body>
""")

def escape_html(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def is_section_heading(text):
    return any(text.startswith(f"{i}.") for i in range(0, 10)) or \
           text in ["Abstract", "Keywords", "References", "Introduction",
                     "Preliminaries", "Related Work",
                     "Experimental Methodology", "Simulation Results",
                     "Real-Data Applications", "Discussion", "Conclusion"] or \
           text.startswith("Contributions")

def is_subsection_heading(text):
    return any(text.startswith(f"{i}.{j}.") for i in range(0, 10) for j in range(0, 10)) or \
           text in ["Parameter Estimation in Classical Statistics",
                     "Tree-Based Methods for Estimation",
                     "Deep Neural Networks",
                     "Statistical Properties of ML Estimators",
                     "ML for Parameter Estimation and Causal Inference",
                     "Uncertainty Quantification for ML Methods",
                     "Simulation Design", "Methods Compared",
                     "Evaluation Protocol",
                     "Linear Scenario", "Semiparametric Scenario",
                     "Nonlinear Scenario", "High-Dimensional Sparse Scenario",
                     "Financial Risk Modeling",
                     "Biomedical Outcome Prediction",
                     "Regimes Where ML Outperforms Traditional Methods",
                     "Computational Trade-offs", "Limitations"]

# Write title
para0 = doc.paragraphs[0].text.strip().replace("\n", " ")
html_parts.append(f'<div class="title">{escape_html(para0)}</div>\n')

# Author info (paragraphs 1-4)
for p in doc.paragraphs[1:5]:
    t = p.text.strip()
    if t:
        html_parts.append(f'<div class="author-info">{escape_html(t)}</div>\n')

# Find table positions in document
table_indices = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip().lower()
    if "table" in t and ("scenario" in t or "simulation" in t or "test" in t or "mse" in t):
        table_indices.append(i)

table_iter = iter(doc.tables)
next_table_caption = 0

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue

    # Check if this paragraph is followed by a table
    # Tables appear at positions 45 (Table 1), 65 (Table 2), 72 (Table 3), 75 (Table 4), 83 (Table 5) in paras
    # We need to check p._element's position relative to tables in the body

    if is_section_heading(t) and t not in ["Abstract", "Keywords", "Contributions"]:
        html_parts.append(f'<div class="section-heading">{escape_html(t)}</div>\n')
    elif is_subsection_heading(t):
        html_parts.append(f'<div class="subsection-heading">{escape_html(t)}</div>\n')
    elif t.startswith("[") and "]" in t[:5]:
        # Reference
        html_parts.append(f'<p class="ref-entry">{escape_html(t)}</p>\n')
    elif t == "Abstract":
        html_parts.append(f'<div class="abstract-label">Abstract</div>\n')
    elif t.startswith("Keywords:"):
        html_parts.append(f'<p class="keywords"><b>Keywords:</b> {escape_html(t[9:])}</p>\n')
    elif t.startswith("Contributions."):
        html_parts.append(f'<p class="body-text"><b>Contributions.</b> {escape_html(t[14:])}</p>\n')
    elif t.startswith("Table"):
        html_parts.append(f'<p class="table-caption">{escape_html(t)}</p>\n')
    else:
        html_parts.append(f'<p class="body-text">{escape_html(t)}</p>\n')

html_parts.append("</body></html>")

with open(HTML_OUT, 'w', encoding='utf-8') as f:
    f.write(''.join(html_parts))

print(f"HTML written: {HTML_OUT}")

# Convert to PDF
HTML(HTML_OUT).write_pdf(PDF_OUT)
size = os.path.getsize(PDF_OUT)
print(f"PDF saved: {PDF_OUT} ({size} bytes)")
